'''
Author: Edward J. C. Ashenbert


'''
from docx import Document
from docx.shared import Inches
import pytesseract


def test_docs_api():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('italic.').italic = True

    p.add_run('bold').bold = True
    p.add_run(' and some ')
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


def test_tesseract_api():
    pytesseract.pytesseract.tesseract_cmd = "C:\Program Files\Tesseract-OCR\\tesseract.exe"

    print(pytesseract.image_to_string(
        "C:\\Users\ASUS\Documents\TestData\Tolerance-Analysis-Images\Images\{0B08A4B0-EC1C-4319-B2A7-ADD3F720F58C}.png",
        'jpn'))


if __name__ == "__main__":
    test_tesseract_api()
